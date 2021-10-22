#!/usr/bin/python
# -*- coding: UTF-8 -*-

from docx import Document
import sqlite3


class seqNode:
    def __init__(self, Questions, Options, Answers):
        self.question = Questions
        self.option = Options
        self.answer = Answers


def getParaQue(doc):
    # 每一段的内容
    i = 0
    queIndex = []
    for para in doc.paragraphs:
        paralist = para.text.split('、')
        try:
            int(paralist[0])
            queIndex.append(i)
        except:
            pass
        i += 1

    return queIndex


def initializeNode(doc, Index):
    docxList = []
    for each in range(len(Index)):  # 通过对索引列表进行索引，实现前一题减后一题的功能，以求出题目的选项、答案和标题
        if each != len(Index) - 1:
            options = []
            for i in range(Index[each] + 1, Index[each + 1]):
                options.append(doc.paragraphs[i].text)

            partialquestion = doc.paragraphs[Index[each]].text.split('、')[1:]  # 取出了第一个数字
            docxNode = seqNode("、".join(partialquestion), options,
                               doc.paragraphs[Index[each]].text.split('【')[1].split('】')[0].split(','))
            docxList.append(docxNode)

    return docxList


def inputDB(docxlist):
    conn = sqlite3.connect('information.db')
    cur = conn.cursor()

    for each in docxlist:  # 将上述结构不重复的写入数据库以供查询
        try:
            sql_text = "insert into quesInfo(Question, Options, Answer) values (" + "'" + str(
                each.question) + "'" + "," + "'" + str("###".join(each.option)) + "'" + "," + "'" + str(
                "###".join(each.answer)) + "'" + ")"

            cur.execute(sql_text)
        except:
            print('重复！跳过！')

    conn.commit()
    cur.close()
    conn.close()


def main():
    doc = Document('./docx/机汽学院安全教育题库.docx')
    inputDB(initializeNode(doc, getParaQue(doc)))


if __name__ == '__main__':
    main()
